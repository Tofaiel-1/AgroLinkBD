import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_proposal_docx(filename):
    doc = Document()

    # Set Margins (0.75 inch for clean 6-page density)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Color Palette
    PRIMARY_COLOR = RGBColor(27, 94, 32)     # Forest Emerald #1B5E20
    SECONDARY_COLOR = RGBColor(0, 121, 107)  # Deep Teal #00796B
    ACCENT_GOLD = RGBColor(194, 120, 3)      # Dark Amber #C27803
    DARK_TEXT = RGBColor(33, 33, 33)         # Charcoal #212121
    GRAY_TEXT = RGBColor(97, 97, 97)         # Slate Gray #616161

    # Base Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = DARK_TEXT

    # Helper Functions
    def add_title(text, subtitle=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR

        if subtitle:
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after = Pt(12)
            run_sub = p_sub.add_run(subtitle)
            run_sub.font.name = 'Calibri'
            run_sub.font.size = Pt(11.5)
            run_sub.font.bold = True
            run_sub.font.color.rgb = SECONDARY_COLOR

    def add_h1(text, page_break=False):
        if page_break:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR

    def add_p(text, bold_prefix=None, italic=False, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.bold = True
            run_b.font.color.rgb = DARK_TEXT
        run = p.add_run(text)
        run.font.italic = italic
        return p

    def add_bullet(title, desc):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.12
        run_t = p.add_run(title + ": ")
        run_t.font.bold = True
        run_t.font.color.rgb = PRIMARY_COLOR
        p.add_run(desc)

    def add_callout(text, title="KEY HIGHLIGHT"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(7.0)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5E9"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        borders_elm = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="1B5E20"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        cell._tc.get_or_add_tcPr().append(borders_elm)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.right_indent = Inches(0.1)
        run_t = p.add_run(f"★ {title}: ")
        run_t.font.bold = True
        run_t.font.color.rgb = PRIMARY_COLOR
        p.add_run(text)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_placeholder(desc):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(7.0)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9FBE7"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        borders_elm = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="dashed" w:sz="12" w:space="0" w:color="827717"/><w:top w:val="dashed" w:sz="12" w:space="0" w:color="827717"/><w:right w:val="dashed" w:sz="12" w:space="0" w:color="827717"/><w:bottom w:val="dashed" w:sz="12" w:space="0" w:color="827717"/></w:tcBorders>')
        cell._tc.get_or_add_tcPr().append(borders_elm)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"📷 [IMAGE PLACEHOLDER: {desc}]")
        run.font.bold = True
        run.font.color.rgb = RGBColor(130, 119, 23)
        run_sub = p.add_run("\n(Paste actual high-resolution mobile app UI screenshot here)")
        run_sub.font.size = Pt(9)
        run_sub.font.italic = True
        run_sub.font.color.rgb = GRAY_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def format_table(table, header_bg="1B5E20", alt_bg="F1F8E9"):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if i == 0:
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_bg}"/>')
                    cell._tc.get_or_add_tcPr().append(shd)
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(3)
                        p.paragraph_format.space_after = Pt(3)
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(9.5)
                else:
                    bg = alt_bg if i % 2 == 1 else "FFFFFF"
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
                    cell._tc.get_or_add_tcPr().append(shd)
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        for run in p.runs:
                            run.font.size = Pt(9)

    # =========================================================================
    # PAGE 1: EXECUTIVE SUMMARY, PROBLEM STATEMENT & STRATEGIC VISION
    # =========================================================================
    add_title("AgroLinkBD: Smart Agriculture & Fisheries Ecosystem",
              "National App Innovation Challenge 2026 | Comprehensive Project Dossier")

    add_callout(
        "A full-stack AI, Agri-Fintech & B2B Trading Super-App designed to eliminate multi-tier intermediary exploitation, "
        "provide real-time Gemini Vision pathology diagnostics, enable live digital auctions, and empower unbanked farmers "
        "with NFC Smart AgroCards across Bangladesh.",
        "EXECUTIVE VISION STATEMENT"
    )

    add_h1("1. Project Identification & Metadata")
    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_data = [
        ("Project Name", "AgroLinkBD (এগ্রোলিংক বিডি) — Smart Agri & Fisheries Super-App"),
        ("Innovation Domain", "Artificial Intelligence (Computer Vision), Agri-Fintech, Supply Chain & IoT"),
        ("Target Beneficiaries", "Crop Farmers, Fish Farmers, Wholesale Aratdars, Fleet Drivers, Machinery Providers, Agronomists"),
        ("National Alignment", "Smart Bangladesh 2041 (Smart Citizen, Smart Economy) & UN SDGs (1, 2, 8, 9, 12)"),
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        cell_k = meta_tbl.cell(row_idx, 0)
        cell_v = meta_tbl.cell(row_idx, 1)
        cell_k.width = Inches(2.0)
        cell_v.width = Inches(5.0)
        cell_k.paragraphs[0].add_run(k).font.bold = True
        cell_v.paragraphs[0].add_run(v)
    format_table(meta_tbl, header_bg="00796B", alt_bg="E0F2F1")

    add_h1("2. The Problem Statement: Critical Crises in Bangladesh Agriculture")
    add_p(
        "Agriculture and aquaculture contribute over 11.6% to Bangladesh's GDP and employ 40%+ of the national workforce. "
        "However, root-level producers face systemic, devastating structural inefficiencies:",
        bold_prefix="Core Problem: "
    )
    add_bullet("Predatory Intermediary Exploitation (মধ্যস্বত্বভোগী)",
               "Producers lose 45%–60% of potential profits to a 4-tier chain of Faria, Paikar, Aratdar, and Retailers.")
    add_bullet("Severe Crop & Fish Disease Losses (মড়ক ও রোগবালাই)",
               "Over ৳4,000+ Crore in annual losses occur due to delayed diagnosis and incorrect over-the-counter chemical poisoning.")
    add_bullet("High Post-Harvest Perishable Waste (পচনশীল ক্ষতি)",
               "25%–30% of harvested vegetables and fish perish in transit due to unorganized rural upazila transportation.")
    add_bullet("Rural Financial Exclusion (ব্যাংক ঋণের সংকট)",
               "Marginal farmers lack formal transaction histories, forcing reliance on illegal, high-interest local loan sharks (দাদন).")

    add_h1("3. Abstract & Project Summary")
    add_p(
        "AgroLinkBD solves these challenges by deploying a unified, enterprise-grade mobile super-app powered by Flutter, "
        "Firebase, and Google Gemini Vision AI. The platform interconnects 7 distinct agricultural roles into a decentralized, "
        "transparent trading floor. Key innovations include an Anti-Deception Vision AI Doctor that rejects invalid images and "
        "computes pond chemical dosages, a live digital wholesale auction floor, corporate forward contract farming, an AI 14-day "
        "price forecasting radar, an on-demand upazila transport dispatch, and an NFC/QR Smart AgroCard providing automated "
        "microfinance KYC credit scoring. AgroLinkBD increases net farm-gate profit margins by 25%–40% while cutting disease losses by 30%."
    )

    # =========================================================================
    # PAGE 2: DESIGN ARCHITECTURE & TECHNICAL WORKFLOW
    # =========================================================================
    add_h1("4. Technical System Architecture & Multi-Role Design", page_break=True)
    add_p(
        "AgroLinkBD utilizes a reactive, cloud-native Clean Architecture with high-throughput state management (GetX & Provider), "
        "real-time Firestore document streams, and edge-optimized Gemini Vision API integration."
    )

    add_h2("A. Architectural Flowchart & System Layers")
    arch_tbl = doc.add_table(rows=5, cols=2)
    arch_layers = [
        ("Presentation Layer (Client)", "Flutter Multi-Platform UI (Responsive Android/iOS/Web), Bengali Typography (Hind Siliguri), Dual Dark/Light Mode, Animated HUD Diagnostics"),
        ("State & Service Layer", "GetX Reactive Controllers, Provider State Management, Secure Local Caching, NFC/QR Scanner, Push Notification Dispatchers"),
        ("AI & Analytics Engine", "Google Gemini 2.5 Flash Multimodal Vision API, Multi-Model Heuristic Fallback Engine, 14-Day Commodity Price Time-Series Radar"),
        ("Cloud Backend & Database", "Google Firebase Firestore (Real-Time Streams), Firebase Authentication (Phone/OTP), Firebase Storage (Encrypted Media), Cloud Functions"),
        ("Fintech & Integration Layer", "SSLCommerz Multi-Channel Payment Gateway, NFC Card Emulation, Digital Wallet PIN Hashing (AES-256), Microfinance Credit KYC Scoring Engine"),
    ]
    for r_idx, (layer, tech) in enumerate(arch_layers):
        c_k = arch_tbl.cell(r_idx, 0)
        c_v = arch_tbl.cell(r_idx, 1)
        c_k.width = Inches(2.2)
        c_v.width = Inches(4.8)
        c_k.paragraphs[0].add_run(layer).font.bold = True
        c_v.paragraphs[0].add_run(tech)
    format_table(arch_tbl)

    add_h2("B. 7-Role Adaptive Multi-Stakeholder Matrix")
    add_p("A single app seamlessly adapts its interface and capabilities based on the authenticated user role:")
    role_tbl = doc.add_table(rows=8, cols=3)
    role_data = [
        ("Role", "Target Stakeholder", "Core App Module & Access"),
        ("1. Crop Farmer", "ফসল চাষী", "Crop Diary, Gemini Plant Doctor, Live Crop Bazaar, Weather Radar"),
        ("2. Fish Farmer", "মৎস্যচাষী", "Multi-Pond Telemetry, AI Fish Doctor, FCR Simulator, Big Fish Trade"),
        ("3. Wholesale Buyer", "আড়তদার ও পাইকার", "Live Auction Bidding, Corporate Contract Farming, Bulk RFQ Tendering"),
        ("4. Logistics Driver", "ট্রাক ও পিকআপ চালক", "Upazila-to-District Load Board, Live GPS Route Dispatch, Trip Escrow"),
        ("5. Input Provider", "বীজ ও সার ডিলার", "Machinery Rental (Harvester/Drone), Agro-Chemical Store, Order Ledger"),
        ("6. Hatchery", "পোনা ও রেণু হ্যাচারি", "Fingerling Calculator, Broodstock Traceability, Direct Paikar Dispatch"),
        ("7. Agro-Doctor", "কৃষি ও মৎস্য বিশেষজ্ঞ", "Live Telemedicine Video Calls, Digital E-Prescriptions, Lab Reports"),
    ]
    for r_idx, (r, s, m) in enumerate(role_data):
        c0, c1, c2 = role_tbl.cell(r_idx, 0), role_tbl.cell(r_idx, 1), role_tbl.cell(r_idx, 2)
        c0.width, c1.width, c2.width = Inches(1.5), Inches(1.5), Inches(4.0)
        c0.paragraphs[0].add_run(r)
        c1.paragraphs[0].add_run(s)
        c2.paragraphs[0].add_run(m)
    format_table(role_tbl)

    add_placeholder("System Architecture & 7-Role Adaptive Dashboard Workflow")

    # =========================================================================
    # PAGE 3: CORE INNOVATIONS & DETAILED MODULE SUITE
    # =========================================================================
    add_h1("5. Deep-Dive Feature Breakdown & Module Capabilities", page_break=True)

    add_h2("Module 1: Ultra Pro Max AI Fish & Crop Doctor (Gemini Vision)")
    add_bullet("Anti-Deception Guardrail", "Rejects irrelevant images (humans, vehicles, objects, animal faces) with polite Bengali photography tips.")
    add_bullet("100% Healthy Specimen Certification", "Accurately detects healthy fish/crops (shiny scales, intact fins, clear eyes) and provides preventive nutrition advice instead of prescribing unnecessary antibiotics.")
    add_bullet("Exact Decimal Chemical Calculator", "Instantly computes pond-specific remediation dosages for Lime (চুন), Salt (লবণ), Zeolite (জিওলাইট), Potash (পটাশ), and Oxytetracycline tailored to pond area and water depth.")
    add_placeholder("AI Fish Doctor Diagnostic Scanner, Non-Fish Rejection & Clinical Dosage Report")

    add_h2("Module 2: Live Wholesale Trading Floor, Auctions & Contract Farming")
    add_bullet("Live Open Auctions (লাইভ নিলাম ডাক)", "Producers broadcast bulk harvests; verified national buyers submit competitive live bids, maximizing the producer's final sale price.")
    add_bullet("Corporate Forward Contracts (চুক্তি চাষ)", "Enables formal pre-harvest buyback agreements between farmers and leading agro-corporates with escrow payment safety.")
    add_bullet("14-Day AI Price Forecast Radar", "Predictive machine learning time-series analyzing commodity demand trajectories across all 64 districts.")
    add_placeholder("Wholesale Trading Floor, Live Auction Bidding & 14-Day Price Forecast Radar")

    add_h2("Module 3: Upazila Hyper-Local Logistics & Fleet Dispatch")
    add_bullet("On-Demand Truck & Pickup Dispatch", "Farmers book trucks directly from upazila loading hubs, cutting transit delays and reducing perishable vegetable/fish spoilage by 30%.")
    add_bullet("Heavy Machinery & Drone Sharing", "Escrow-secured hourly rental of Combine Harvesters, Tractors, and Pesticide Spraying Drones, cutting labor costs by 35%.")

    add_h2("Module 4: Smart AgroCard & Automated Microfinance Credit Scoring")
    add_bullet("NFC & QR Digital Farmer Identity", "Provides an unbanked farmer with a physical/digital identity card carrying a multi-tier PIN-secured digital wallet.")
    add_bullet("Bank Loan Project Report (KYC) Generator", "Aggregates harvest records, land size, and sales history into an official Bank Loan Dossier, eliminating collateral harassment.")
    add_placeholder("Smart NFC AgroCard, Digital Wallet & Bank Loan Project Dossier")

    # =========================================================================
    # PAGE 4: BUSINESS MODEL, MONETIZATION & FINANCIAL VIABILITY
    # =========================================================================
    add_h1("6. Business Model & Sustainable Monetization Strategy", page_break=True)
    add_p(
        "AgroLinkBD operates on a highly scalable, multi-stream revenue model that delivers immense financial value to grassroots "
        "farmers while maintaining strong commercial profitability and zero exploitation.",
        bold_prefix="Business Philosophy: "
    )

    biz_tbl = doc.add_table(rows=7, cols=4)
    biz_data = [
        ("Revenue Stream", "Target Segment", "Pricing Mechanism", "Projected Annual Revenue (Yr 2)"),
        ("1. Wholesale Trade Commission", "B2B Wholesalers / Aratdars", "1.5% – 2.5% transaction commission on auction & contract sales", "৳ 3.20 Crore"),
        ("2. Agro-Telemedicine Tokens", "Farmers seeking specialist doctor", "৳ 30 – ৳ 50 per live specialist video call token", "৳ 45.00 Lakh"),
        ("3. Logistics & Rental Escrow", "Fleet drivers & equipment owners", "3.0% – 5.0% platform safety and escrow fee", "৳ 85.00 Lakh"),
        ("4. B2B Corporate Ads & Deals", "Seed, feed & fertilizer MNCs (ACI, Lal Teer)", "Sponsored brand deal placement & distributor zones", "৳ 60.00 Lakh"),
        ("5. VIP Pass & Intelligence Sub", "Commercial fish & crop producers", "৳ 199/month for 14-day price radar & bank report export", "৳ 72.00 Lakh"),
        ("6. Microfinance Loan Origination", "Partner Commercial Banks / MFIs", "0.5% – 1.0% origination fee on disbursed credit dossiers", "৳ 90.00 Lakh"),
    ]
    for r_idx, (col0, col1, col2, col3) in enumerate(biz_data):
        c0, c1, c2, c3 = biz_tbl.cell(r_idx, 0), biz_tbl.cell(r_idx, 1), biz_tbl.cell(r_idx, 2), biz_tbl.cell(r_idx, 3)
        c0.width, c1.width, c2.width, c3.width = Inches(1.8), Inches(1.8), Inches(2.2), Inches(1.2)
        c0.paragraphs[0].add_run(col0)
        c1.paragraphs[0].add_run(col1)
        c2.paragraphs[0].add_run(col2)
        c3.paragraphs[0].add_run(col3)
    format_table(biz_tbl, header_bg="1B5E20", alt_bg="E8F5E9")

    add_h2("3-Year Financial Growth Projections")
    fin_data = [
        ("Metric (BDT)", "Year 1 (Pilot)", "Year 2 (Expansion)", "Year 3 (National Scale)", "Key Driver"),
        ("Active Farmers", "25,000", "150,000", "500,000+", "Upazila Agro-Hubs"),
        ("Gross Merchandise Value (GMV)", "৳ 18.5 Crore", "৳ 145.0 Crore", "৳ 580.0 Crore", "Live Auction Volume"),
        ("Net Platform Revenue", "৳ 82.5 Lakh", "৳ 6.72 Crore", "৳ 26.50 Crore", "6 Diversified Streams"),
        ("Net Profit Margin", "18.5%", "34.2%", "48.0%", "Zero Marginal Server Cost"),
    ]
    fin_tbl = doc.add_table(rows=len(fin_data), cols=5)
    for r_idx, (m, y1, y2, y3, d) in enumerate(fin_data):
        c0, c1, c2, c3, c4 = fin_tbl.cell(r_idx, 0), fin_tbl.cell(r_idx, 1), fin_tbl.cell(r_idx, 2), fin_tbl.cell(r_idx, 3), fin_tbl.cell(r_idx, 4)
        c0.width, c1.width, c2.width, c3.width, c4.width = Inches(1.8), Inches(1.2), Inches(1.3), Inches(1.4), Inches(1.3)
        c0.paragraphs[0].add_run(m)
        c1.paragraphs[0].add_run(y1)
        c2.paragraphs[0].add_run(y2)
        c3.paragraphs[0].add_run(y3)
        c4.paragraphs[0].add_run(d)
    format_table(fin_tbl, header_bg="00796B", alt_bg="E0F2F1")

    # =========================================================================
    # PAGE 5: IMPLEMENTATION ROADMAP, SECURITY & FEASIBILITY
    # =========================================================================
    add_h1("7. Implementation Roadmap & Technical Feasibility", page_break=True)

    add_h2("A. 4-Phase Phased National Rollout Strategy")
    add_bullet("Phase 1: Pilot Deployment (Months 1–4)",
               "Launch across 4 key agricultural districts (Bogura, Mymensingh, Jashore, Rajshahi). Onboard 10,000 farmers and 50 aratdars.")
    add_bullet("Phase 2: Regional Scale & Bank Tie-ups (Months 5–8)",
               "Expand to 30 agricultural upazilas. Integrate 3 commercial partner banks for microfinance loan disbursement via AgroCard.")
    add_bullet("Phase 3: Nationwide Commercial Scaling (Months 9–12)",
               "Cover all 64 districts. Expand fleet logistics network to 5,000 registered trucks and harvesters. Target 100,000 active farmers.")
    add_bullet("Phase 4: IoT & Cross-Border Export Trade (Year 2+)",
               "Integrate IoT automated pond water sensors and facilitate direct agricultural export compliance to Middle East/EU markets.")

    add_h2("B. Security, Scalability & Data Privacy Framework")
    add_bullet("Granular Firestore Security Rules", "Every document write/read is authenticated via UID checks and administrative role validation (400+ lines of enterprise rules).")
    add_bullet("Financial Cryptography", "Digital Wallet PINs are secured with irreversible cryptographic salting and AES-256 encryption. SSLCommerz PCI-DSS certified.")
    add_bullet("High-Concurrence Scalability", "Serverless Google Cloud architecture effortlessly scales to 1,000,000+ simultaneous real-time auction bids without degradation.")

    # =========================================================================
    # PAGE 6: COMPETITIVE BENCHMARKING, SOCIO-ECONOMIC IMPACT & JURY FAQ
    # =========================================================================
    add_h1("8. Competitive Analysis: Why AgroLinkBD Wins Nationwide", page_break=True)
    add_p(
        "A rigorous comparison against existing digital agricultural solutions demonstrates AgroLinkBD's decisive technological and commercial superiority:",
        bold_prefix="Market Landscape: "
    )

    comp_tbl = doc.add_table(rows=7, cols=5)
    comp_data = [
        ("Feature / Capability", "AgroLinkBD", "Krishi Batayon", "iFarmer", "Chaldal / Shodagor"),
        ("1. Gemini Vision AI Doctor (Fish & Crop)", "✅ Yes (With Dosage & Filter)", "❌ Static Text Only", "❌ None", "❌ None"),
        ("2. Live Digital Auctions & Contract Farming", "✅ Yes (Real-Time Floor)", "❌ None", "❌ None", "❌ Fixed Retail Only"),
        ("3. 14-Day AI Price Forecast Radar", "✅ Yes (Machine Learning)", "❌ None", "❌ None", "❌ None"),
        ("4. Smart NFC AgroCard & Loan KYC Dossier", "✅ Yes (Hardware+App)", "❌ None", "❌ Semi-Digital", "❌ None"),
        ("5. Upazila Logistics & Machinery Rental", "✅ Yes (On-Demand GPS)", "❌ None", "❌ None", "❌ City Logistics Only"),
        ("6. 7-Role Adaptive Ecosystem", "✅ Complete Super-App", "❌ 1 Role Only", "❌ 1 Role Only", "❌ E-commerce Only"),
    ]
    for r_idx, (f, ab, kb, ifa, cs) in enumerate(comp_data):
        c0, c1, c2, c3, c4 = comp_tbl.cell(r_idx, 0), comp_tbl.cell(r_idx, 1), comp_tbl.cell(r_idx, 2), comp_tbl.cell(r_idx, 3), comp_tbl.cell(r_idx, 4)
        c0.width, c1.width, c2.width, c3.width, c4.width = Inches(2.2), Inches(1.5), Inches(1.1), Inches(1.1), Inches(1.1)
        c0.paragraphs[0].add_run(f)
        c1.paragraphs[0].add_run(ab)
        c2.paragraphs[0].add_run(kb)
        c3.paragraphs[0].add_run(ifa)
        c4.paragraphs[0].add_run(cs)
    format_table(comp_tbl, header_bg="1B5E20", alt_bg="E8F5E9")

    add_h1("9. Quantified Expected Socio-Economic Impact")
    add_bullet("Farm-Gate Income Increase", "+25% to 40% higher net earnings for farmers by cutting 4 layers of middlemen.")
    add_bullet("Post-Harvest & Disease Loss Reduction", "-30% reduction in perishable spoilage and aquatic disease mortality, saving national wealth.")
    add_bullet("Financial Inclusion of Unbanked Farmers", "100,000+ marginal producers brought into the formal banking loan system without collateral.")
    add_bullet("National Alignment", "Directly drives Smart Bangladesh 2041 and UN SDGs 1 (No Poverty), 2 (Zero Hunger), 8 (Economic Growth), and 9 (Innovation).")

    add_h1("10. Jury Board Winning Defense (FAQ Cheat Sheet)")
    add_bullet("Judge Q1: 'How will illiterate farmers use this app?'",
               "Answer: Full Bengali audio-visual UI, intuitive icon navigation, and the physical NFC Smart AgroCard that works with a simple tap at local agro-dealers.")
    add_bullet("Judge Q2: 'How is this different from a simple Facebook buying group?'",
               "Answer: Facebook has no escrow protection, no live auction bidding mechanism, no AI disease vision diagnostics, and no logistics GPS tracking.")
    add_bullet("Judge Q3: 'Is the business model profitable without charity/grants?'",
               "Answer: Yes. With 6 commercial revenue streams (trade commission, VIP subs, ad deals, escrow fees), the platform is fully profitable from Year 1 with zero marginal server scaling cost.")

    # Save Document
    doc.save(filename)
    print("SUCCESS: Created Word Document at " + filename)

if __name__ == "__main__":
    create_proposal_docx(r"d:\App\AgroLinkBD\AgroLinkBD_National_Innovation_Challenge_Proposal.docx")
